import io
import json
import re
from pathlib import Path
from flask import Blueprint, request, jsonify
from flask_jwt_extended import jwt_required, get_jwt_identity
from werkzeug.utils import secure_filename

from backend.services.rag_service import refresh_retriever
from backend.database.repository import DocumentRepository
from backend.database.models import User
from backend.utils.rbac import require_permission, get_current_organization_id
from backend.utils.audit import log_audit
from backend.utils.abac import evaluate_abac_policies


def _check_clearance_access(user, document):
    """Return (allowed: bool, reason: str|None). Admins (caller must check) bypass this."""
    if document.clearance_level_id is None:
        return True, None
    if user.clearance_level_id is None:
        return False, (
            f"Document requires clearance '{document.clearance_level.name}' "
            "but your account has no clearance level assigned"
        )
    if user.clearance_level.level >= document.clearance_level.level:
        return True, None
    return False, (
        f"Insufficient clearance: document requires level {document.clearance_level.level} "
        f"({document.clearance_level.name}), you have level {user.clearance_level.level} "
        f"({user.clearance_level.name})"
    )

documents_bp = Blueprint('documents', __name__, url_prefix='/api')

# Path configuration
BASE_DIR = Path(__file__).resolve().parents[1]
DOCS_PATH = BASE_DIR / "data" / "docs"
HIGH_DOCS_PATH = BASE_DIR / "data" / "docs" / "high"
CONVERTED_PATH = BASE_DIR / "data" / "docs_converted"
METADATA_PATH = BASE_DIR / "data" / "docs_metadata.json"

# Accept all file types - no restrictions
ALLOWED_EXTENSIONS = None  # Accept everything

# Ensure directories exist
DOCS_PATH.mkdir(parents=True, exist_ok=True)
CONVERTED_PATH.mkdir(parents=True, exist_ok=True)

def allowed_file(filename):
    # Accept all files
    return True

def load_metadata():
    """Load document metadata from JSON file"""
    if METADATA_PATH.exists():
        with open(METADATA_PATH, 'r') as f:
            return json.load(f)
    return {}

def save_metadata(metadata):
    """Save document metadata to JSON file"""
    with open(METADATA_PATH, 'w') as f:
        json.dump(metadata, f, indent=2)

def extract_text_from_file(file_path):
    """Extract text content from various file formats"""
    file_extension = file_path.suffix.lower()

    try:
        # Plain text files
        if file_extension in ['.txt', '.md', '.log', '.csv', '.json', '.xml', '.html', '.htm']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        # PDF files
        elif file_extension == '.pdf':
            try:
                import pdfplumber
                import re
                pages = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
                        pages.append(page_text)
                raw = '\n'.join(pages)
                # Fix hyphenated line-breaks: "re-\ntrieve" → "retrieve"
                raw = re.sub(r'-\n(\s*)', '', raw)
                # Collapse multiple blank lines
                raw = re.sub(r'\n{3,}', '\n\n', raw)
                # Join mid-sentence line breaks (line doesn't end with punctuation)
                raw = re.sub(r'(?<![.!?:])\n(?=[a-z])', ' ', raw)
                return raw.strip()
            except ImportError:
                pass
            try:
                import PyPDF2
                import re
                pages = []
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        pages.append(page.extract_text() or "")
                raw = '\n'.join(pages)
                raw = re.sub(r'-\n(\s*)', '', raw)
                raw = re.sub(r'\n{3,}', '\n\n', raw)
                raw = re.sub(r'(?<![.!?:])\n(?=[a-z])', ' ', raw)
                return raw.strip()
            except ImportError:
                with open(file_path, 'rb') as f:
                    return f.read().decode('utf-8', errors='ignore')

        # DOCX files
        elif file_extension == '.docx':
            try:
                import docx
                doc = docx.Document(file_path)
                return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                # Fallback: read as text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()

        # DOC files (older Word format)
        elif file_extension == '.doc':
            try:
                import pypandoc
                return pypandoc.convert_file(str(file_path), 'plain')
            except (ImportError, RuntimeError):
                # Fallback: read as text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()

        # Other files - try reading as text
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        # Last resort: read as binary and decode
        try:
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore')
        except:
            return f"[Unable to extract text from {file_path.name}]"

def extract_text_from_bytes(raw_bytes: bytes, filename: str) -> str:
    """Extract text from bytes in memory — never writes plaintext to disk."""
    import io as _io
    ext = Path(filename).suffix.lower()
    try:
        if ext == '.pdf':
            try:
                import pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    pages = [p.extract_text(x_tolerance=2, y_tolerance=2) or "" for p in pdf.pages]
                raw = '\n'.join(pages)
                raw = re.sub(r'-\n(\s*)', '', raw)
                raw = re.sub(r'\n{3,}', '\n\n', raw)
                raw = re.sub(r'(?<![.!?:])\n(?=[a-z])', ' ', raw)
                return raw.strip()
            except ImportError:
                import PyPDF2
                reader = PyPDF2.PdfReader(_io.BytesIO(raw_bytes))
                pages = [p.extract_text() or "" for p in reader.pages]
                raw = '\n'.join(pages)
                raw = re.sub(r'-\n(\s*)', '', raw)
                raw = re.sub(r'\n{3,}', '\n\n', raw)
                raw = re.sub(r'(?<![.!?:])\n(?=[a-z])', ' ', raw)
                return raw.strip()
        elif ext == '.docx':
            import docx
            doc = docx.Document(_io.BytesIO(raw_bytes))
            return '\n'.join(p.text for p in doc.paragraphs)
        else:
            return raw_bytes.decode('utf-8', errors='ignore')
    except Exception:
        return raw_bytes.decode('utf-8', errors='ignore')


def convert_to_txt(original_path, converted_path):
    """Convert any file to .txt format for RAG system"""
    text_content = extract_text_from_file(original_path)

    # Save as .txt file
    txt_filename = original_path.stem + '.txt'
    txt_path = converted_path / txt_filename

    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(text_content)

    return txt_path, text_content

def scan_document(filename, content):
    """Basic security scanning for documents"""
    issues = []

    # Check for PII patterns
    email_pattern = r'\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b'
    phone_pattern = r'\b\d{10,15}\b'

    if re.search(email_pattern, content, re.IGNORECASE):
        issues.append("PII: Email detected")
    if re.search(phone_pattern, content):
        issues.append("PII: Phone number detected")

    # Check for potential injection patterns
    injection_keywords = ['<script>', 'javascript:', 'onerror=', 'eval(', 'exec(']
    for keyword in injection_keywords:
        if keyword.lower() in content.lower():
            issues.append(f"Injection: {keyword} detected")

    return issues

def _ingest_high_sensitivity(raw_bytes: bytes, filename: str, doc_id: str, text_content: str) -> None:
    """
    Pre-storage encrypted ingest pipeline for HIGH sensitivity documents.

    Plaintext NEVER touches disk:
    1. Encrypt raw_bytes in RAM → write ONLY ciphertext to data/docs/high/
    2. Ingest text chunks into SecureRAGProvider (Chroma chunk_ids + enc_store ciphertext)
    """
    from flask import current_app
    from langchain_core.documents import Document as LCDocument
    from backend.api.secure_query import _get_secure_provider
    from backend.security.sag_crypto import encrypt, derive_key
    import logging

    _log = logging.getLogger("zerosec.documents")

    try:
        HIGH_DOCS_PATH.mkdir(parents=True, exist_ok=True)
        high_file_path = HIGH_DOCS_PATH / filename

        # 1. Encrypt raw bytes in RAM — write ONLY ciphertext to disk
        secret_key = current_app.config.get('SECRET_KEY', '')
        key_version = current_app.config.get('SAG_KEY_VERSION', 1)
        key = derive_key(secret_key, doc_id, 'HIGH', key_version)
        ciphertext, nonce, auth_tag = encrypt(raw_bytes, key)
        high_file_path.write_bytes(nonce + auth_tag + ciphertext)

        # 2. Ingest text chunks into encrypted pipeline
        file_ext = Path(filename).suffix
        provider = _get_secure_provider()
        doc = LCDocument(
            page_content=text_content,
            metadata={
                'source':    str(high_file_path),
                'filename':  filename,
                'file_type': file_ext,
                'doc_id':    doc_id,
            }
        )
        provider.ingest_documents([doc])

        _log.info("[HIGH ingest] %s encrypted and ingested (plaintext never written to disk)", filename)

    except Exception as exc:
        _log.error("[HIGH ingest] failed for %s: %s", filename, exc)
        raise


@documents_bp.route('/documents', methods=['GET'])
@jwt_required()
@require_permission('read')
def get_documents():
    """Get all documents with metadata"""
    try:
        organization_id = get_current_organization_id()
        user_id = get_jwt_identity()

        # Get documents from database
        db_documents = DocumentRepository.get_all_documents(organization_id)

        from flask_jwt_extended import get_jwt
        is_admin = 'admin' in get_jwt().get('permissions', [])
        user = None if is_admin else User.query.get(int(user_id))

        documents = []
        for doc in db_documents:
            if not is_admin:
                allowed, _ = _check_clearance_access(user, doc)
                if not allowed:
                    continue

            # Try to get file stats — HIGH docs live in HIGH_DOCS_PATH (stored as ciphertext)
            if doc.sensitivity == 'High':
                file_path = HIGH_DOCS_PATH / doc.filename
            else:
                file_path = DOCS_PATH / doc.filename
            file_size = file_path.stat().st_size if file_path.exists() else 0

            documents.append({
                'id': doc.document_id,
                'name': doc.filename,
                'sensitivity': doc.sensitivity,
                'clearance_level': doc.clearance_level.name if doc.clearance_level else None,
                'size': file_size,
                'uploaded_at': doc.created_at.isoformat()
            })

        # Log audit event
        log_audit(
            organization_id=organization_id,
            user_id=user_id,
            action='documents_listed',
            target_type='Document'
        )

        return jsonify({'documents': documents}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@documents_bp.route('/documents/upload', methods=['POST'])
@jwt_required()
@require_permission('create')
def upload_document():
    """Upload a new document"""
    try:
        organization_id = get_current_organization_id()
        user_id = get_jwt_identity()

        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Secure the filename
        filename = secure_filename(file.filename)

        # Check if document already exists in database
        existing_doc = DocumentRepository.get_document_by_filename(organization_id, filename)
        if existing_doc:
            return jsonify({'error': 'Document already exists'}), 409

        # Read file into memory — HIGH sensitivity files must NEVER touch disk as plaintext
        raw_bytes = file.read()

        # Extract text content for scanning (from memory)
        try:
            text_content = extract_text_from_bytes(raw_bytes, filename)
        except Exception as e:
            text_content = raw_bytes.decode('utf-8', errors='ignore')

        # Scan document for issues
        issues = scan_document(filename, text_content)

        # Determine sensitivity
        provided_sensitivity = request.form.get('sensitivity', '').lower()
        if provided_sensitivity in ['high', 'medium', 'low']:
            sensitivity = provided_sensitivity.capitalize()
        else:
            sensitivity = 'High' if any('PII' in issue for issue in issues) else 'Medium' if issues else 'Low'

        # Optional clearance level assignment on upload
        raw_clearance = request.form.get('clearance_level_id')
        clearance_level_id = int(raw_clearance) if raw_clearance and raw_clearance.isdigit() else None

        # Route based on sensitivity — HIGH never writes plaintext to disk
        if sensitivity == 'High':
            # Create DB record before ingest (needed for doc_id)
            document = DocumentRepository.create_document(
                organization_id=organization_id,
                filename=filename,
                storage_ref=str(HIGH_DOCS_PATH / filename),
                sensitivity=sensitivity,
                clearance_level_id=clearance_level_id,
                user_id=user_id
            )
            doc_id_str = str(document.document_id)

            # Run ingest in a background thread — embedding can take 30+ seconds
            # The document record is already saved; ingest failure is logged but
            # does not block the upload response.
            import threading
            from flask import current_app

            app_ctx = current_app._get_current_object()

            def _bg_ingest(app, raw, fname, did, text):
                with app.app_context():
                    try:
                        _ingest_high_sensitivity(raw, fname, did, text)
                    except Exception as exc:
                        import logging
                        logging.getLogger("zerosec.documents").error(
                            "[HIGH upload] background ingest failed for %s: %s", fname, exc
                        )

            threading.Thread(
                target=_bg_ingest,
                args=(app_ctx, raw_bytes, filename, doc_id_str, text_content),
                daemon=True,
            ).start()
        else:
            # Standard pipeline — save plaintext to disk, refresh Chroma vectorstore
            file_path = DOCS_PATH / filename
            with open(str(file_path), 'wb') as f:
                f.write(raw_bytes)

            document = DocumentRepository.create_document(
                organization_id=organization_id,
                filename=filename,
                storage_ref=str(file_path),
                sensitivity=sensitivity,
                clearance_level_id=clearance_level_id,
                user_id=user_id
            )
            try:
                refresh_retriever()
            except Exception:
                pass

        # Log audit event
        log_audit(
            organization_id=organization_id,
            user_id=user_id,
            action='document_uploaded',
            target_type='Document',
            target_id=document.document_id,
            metadata={'filename': filename, 'sensitivity': sensitivity, 'issues_count': len(issues)}
        )

        return jsonify({
            'message': 'File uploaded successfully',
            'document': {
                'id': document.document_id,
                'name': filename,
                'sensitivity': sensitivity,
                'issues_count': len(issues)
            }
        }), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@documents_bp.route('/documents/<int:document_id>', methods=['GET'])
@jwt_required()
@require_permission('read')
def get_document_details(document_id):
    """Get detailed information about a specific document"""
    try:
        organization_id = get_current_organization_id()
        user_id = get_jwt_identity()

        # Get document from database
        document = DocumentRepository.get_document_by_id(document_id)

        if not document or document.organization_id != organization_id:
            return jsonify({'error': 'Document not found'}), 404

        from flask_jwt_extended import get_jwt
        if 'admin' not in get_jwt().get('permissions', []):
            user = User.query.get(int(user_id))
            allowed, reason = _check_clearance_access(user, document)
            if not allowed:
                log_audit(organization_id=organization_id, user_id=user_id,
                          action='clearance_access_denied', target_type='Document',
                          target_id=document_id, metadata={'reason': reason})
                return jsonify({'error': 'Access denied', 'reason': reason}), 403
            abac_ok, abac_reason = evaluate_abac_policies(user, document, organization_id)
            if not abac_ok:
                log_audit(organization_id=organization_id, user_id=user_id,
                          action='abac_access_denied', target_type='Document',
                          target_id=document_id, metadata={'reason': abac_reason})
                return jsonify({'error': 'Access denied by policy', 'reason': abac_reason}), 403

        file_path = DOCS_PATH / document.filename

        # Extract content preview if file exists
        content_preview = ""
        file_size = 0
        if file_path.exists():
            file_size = file_path.stat().st_size
            try:
                text_content = extract_text_from_file(file_path)
                content_preview = text_content[:500] + "..." if len(text_content) > 500 else text_content
            except Exception:
                content_preview = "[Unable to extract content preview]"

        # Get activity log from audit
        from backend.utils.audit import get_document_activity
        activity_log = get_document_activity(document_id, limit=10)

        # Build document details
        doc_details = {
            'id': document.document_id,
            'name': document.filename,
            'sensitivity': document.sensitivity,
            'clearance_level': document.clearance_level.name if document.clearance_level else None,
            'size': file_size,
            'uploaded_at': document.created_at.isoformat(),
            'content_preview': content_preview,
            'activity_log': activity_log
        }

        # Log audit event
        log_audit(
            organization_id=organization_id,
            user_id=user_id,
            action='document_viewed',
            target_type='Document',
            target_id=document_id
        )

        return jsonify({'document': doc_details}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@documents_bp.route('/documents/<int:document_id>', methods=['DELETE'])
@jwt_required()
@require_permission('delete')
def delete_document(document_id):
    """Delete a document"""
    try:
        organization_id = get_current_organization_id()
        user_id = get_jwt_identity()

        # Get document from database
        document = DocumentRepository.get_document_by_id(document_id)

        if not document or document.organization_id != organization_id:
            return jsonify({'error': 'Document not found'}), 404

        from flask_jwt_extended import get_jwt
        if 'admin' not in get_jwt().get('permissions', []):
            user = User.query.get(int(user_id))
            allowed, reason = _check_clearance_access(user, document)
            if not allowed:
                log_audit(organization_id=organization_id, user_id=user_id,
                          action='clearance_access_denied', target_type='Document',
                          target_id=document_id, metadata={'reason': reason})
                return jsonify({'error': 'Access denied', 'reason': reason}), 403

        filename = document.filename

        # Delete physical file — check both standard and high/ paths
        for candidate in [DOCS_PATH / filename, HIGH_DOCS_PATH / filename]:
            if candidate.exists():
                candidate.unlink()

        # If HIGH sensitivity, remove from encrypted pipeline
        if document.sensitivity == 'High':
            try:
                from backend.api.secure_query import _get_secure_provider
                _get_secure_provider().delete_documents([filename])
            except Exception as exc:
                import logging
                logging.getLogger("zerosec.documents").warning(
                    "[delete] encrypted pipeline cleanup failed for %s: %s", filename, exc
                )

        # Delete from database (cascades to uploads, canary tokens, etc.)
        DocumentRepository.delete_document(document_id)

        # Log audit event
        log_audit(
            organization_id=organization_id,
            user_id=user_id,
            action='document_deleted',
            target_type='Document',
            target_id=document_id,
            metadata={'filename': filename}
        )

        # Auto-refresh vectorstore
        try:
            refresh_retriever()
        except Exception:
            pass

        return jsonify({'message': 'Document deleted successfully'}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@documents_bp.route('/documents/<int:document_id>/vault-inspect', methods=['GET'])
@jwt_required()
@require_permission('read')
def vault_inspect(document_id):
    """
    Return raw encrypted metadata for a HIGH sensitivity document.
    Never returns plaintext — only ciphertext hex, nonce, auth_tag, algorithm info.
    Used by the Vault Inspector UI to prove data is encrypted at rest.
    """
    try:
        organization_id = get_current_organization_id()
        user_id = get_jwt_identity()

        document = DocumentRepository.get_document_by_id(document_id)
        if not document or document.organization_id != organization_id:
            return jsonify({'error': 'Document not found'}), 404

        if document.sensitivity != 'High':
            return jsonify({'error': 'Vault Inspector is only available for HIGH sensitivity documents'}), 400

        from backend.database import enc_store as _enc_store
        rows = _enc_store.get_all_rows()
        doc_rows = [r for r in rows if r['filename'] == document.filename]

        if not doc_rows:
            return jsonify({'error': 'No encrypted chunks found — document may not be ingested yet'}), 404

        chunks = []
        for r in doc_rows:
            ct = bytes(r['ciphertext'])
            nonce = bytes(r['nonce'])
            auth_tag = bytes(r['auth_tag'])
            chunks.append({
                'chunk_id':     r['chunk_id'][:16] + '...',
                'nonce_hex':    nonce.hex(),
                'auth_tag_hex': auth_tag.hex(),
                'ciphertext_hex': ct[:32].hex() + '...',
                'ciphertext_bytes': len(ct),
                'key_version':  r['key_version'],
                'created_at':   r['created_at'],
            })

        log_audit(
            organization_id=organization_id,
            user_id=user_id,
            action='vault_inspect',
            target_type='Document',
            target_id=document_id,
            metadata={'filename': document.filename, 'chunks_inspected': len(chunks)},
        )

        return jsonify({
            'filename':    document.filename,
            'sensitivity': document.sensitivity,
            'algorithm':   'AES-256-GCM',
            'kdf':         'HKDF-SHA256',
            'chunks':      chunks,
            'total_chunks': len(chunks),
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@documents_bp.route('/documents/refresh', methods=['POST'])
def refresh_documents():
    """Force refresh the RAG retriever to pick up new documents"""
    try:
        refresh_retriever()
        return jsonify({'message': 'RAG retriever refreshed successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500
