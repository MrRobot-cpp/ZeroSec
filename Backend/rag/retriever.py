from pathlib import Path
from hashlib import md5
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_core.documents import Document
from backend.rag.chunker import chunk_documents as _chunk_documents

# -------------------------
# CONFIG
# -------------------------
BASE_DIR = Path(__file__).resolve().parents[1]
DOCS_PATH = BASE_DIR / "data" / "docs"
PERSIST_DIR = BASE_DIR / "data" / "vectorstore"
EMBEDDING_MODEL = "nomic-embed-text"  # Proper embedding model for semantic search

# Chunking config
CHUNK_SIZE = 1000  # Larger chunks = fewer chunks, more context per chunk
CHUNK_OVERLAP = 200  # Higher overlap preserves context at chunk boundaries

# Retriever config — nomic-embed-text produces L2 distances typically in 0.3–1.8 range
TOP_K = 10
DISTANCE_THRESHOLD = 1.8  # allow relevant chunks; real irrelevant ones score > 2.0
MAX_RESULTS = 5

# Global cache
_vectorstore_cache = None
_last_file_hash = None
_embeddings_cache = None


def _get_embeddings():
    """Get cached embeddings instance."""
    global _embeddings_cache
    if _embeddings_cache is None:
        _embeddings_cache = OllamaEmbeddings(model=EMBEDDING_MODEL)
    return _embeddings_cache


def _compute_files_hash():
    """Compute hash of all files (names + sizes + mtimes) for cache invalidation."""
    hash_parts = []
    for f in sorted(DOCS_PATH.glob('*.*')):
        if f.is_file():
            stat = f.stat()
            hash_parts.append(f"{f.name}:{stat.st_size}:{stat.st_mtime}")
    return md5("|".join(hash_parts).encode()).hexdigest()


def _clean_pdf_text(text: str) -> str:
    """Clean PDF-extracted text: remove repeated TOC/slide headers and normalize."""
    import re
    lines = text.split('\n')

    # Detect repeated blocks (slide TOC headers that appear on every page)
    line_counts = {}
    for line in lines:
        stripped = line.strip()
        if len(stripped) > 10:
            line_counts[stripped] = line_counts.get(stripped, 0) + 1

    # Lines appearing 3+ times are repeated headers/footers — remove duplicates
    seen_repeated = set()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if line_counts.get(stripped, 0) >= 3:
            if stripped in seen_repeated:
                continue  # skip duplicate
            seen_repeated.add(stripped)  # keep first occurrence
        cleaned.append(line)

    text = '\n'.join(cleaned)

    # Collapse 3+ consecutive newlines into 2
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Replace common PDF bullet artifacts with clean bullets
    text = text.replace('\uf0d8', '- ')

    return text.strip()


def extract_text_from_file(file_path):
    """Extract text from various file formats."""
    ext = file_path.suffix.lower()

    try:
        if ext in ['.txt', '.md', '.log', '.csv', '.json', '.xml', '.html', '.htm']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        elif ext == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    raw = '\n'.join(page.extract_text() or '' for page in reader.pages)
                    return _clean_pdf_text(raw)
            except Exception:
                return ""

        elif ext == '.docx':
            try:
                import docx
                doc = docx.Document(file_path)
                return '\n'.join(p.text for p in doc.paragraphs)
            except Exception:
                return ""

        elif ext == '.doc':
            try:
                import pypandoc
                return pypandoc.convert_file(str(file_path), 'plain')
            except Exception:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

    except Exception:
        return ""



def load_all_documents():
    """
    Load all documents from docs directory.
    Excludes data/docs/high/ — HIGH sensitivity files are handled by the
    encrypted pipeline (SecureRAGProvider) and must never enter standard Chroma.
    """
    documents = []
    high_dir = DOCS_PATH / "high"

    for file_path in DOCS_PATH.glob('*.*'):
        # Skip anything inside the high/ subdirectory
        if high_dir in file_path.parents:
            continue
        if file_path.is_file():
            text = extract_text_from_file(file_path)
            if text and text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={
                        'source': str(file_path),
                        'filename': file_path.name,
                        'file_type': file_path.suffix
                    }
                ))
    return documents


def _ensure_vectorstore(force_reload=False):
    """
    Ensure vectorstore is built and cached.
    Returns the vectorstore instance.
    """
    global _vectorstore_cache, _last_file_hash

    current_hash = _compute_files_hash()

    # Check if we need to rebuild
    needs_rebuild = (
        force_reload or
        _vectorstore_cache is None or
        current_hash != _last_file_hash
    )

    if needs_rebuild:
        # Load and chunk documents
        documents = load_all_documents()

        if not documents:
            # Create minimal in-memory vectorstore for empty state
            _vectorstore_cache = Chroma.from_texts(
                texts=["No documents available"],
                embedding=_get_embeddings()
            )
        else:
            # Chunk documents for better retrieval
            chunked_docs = _chunk_documents(documents)

            # Create in-memory vectorstore (simpler, avoids file locking issues)
            _vectorstore_cache = Chroma.from_documents(
                documents=chunked_docs,
                embedding=_get_embeddings()
            )

        _last_file_hash = current_hash

    return _vectorstore_cache


def retrieve_with_scores(query: str, force_reload=False):
    """
    Retrieve documents with relevance scores and filter by threshold.
    Only returns documents that are actually relevant to the query.

    Returns: List of (Document, score) tuples where score is normalized 0-1 (higher = more relevant)
    """
    vectorstore = _ensure_vectorstore(force_reload)

    # Get documents with distance scores (lower distance = more similar)
    # Using similarity_search_with_score which returns raw distances
    results = vectorstore.similarity_search_with_score(
        query,
        k=TOP_K
    )

    # Filter by distance threshold - only keep documents within threshold
    # Lower distance = more relevant, so we keep docs with distance < threshold
    filtered_results = [
        (doc, distance) for doc, distance in results
        if distance <= DISTANCE_THRESHOLD
    ]

    # Convert distance to similarity score (0-1, higher = more similar)
    # Using formula: similarity = 1 / (1 + distance)
    scored_results = [
        (doc, round(1 / (1 + distance), 3))
        for doc, distance in filtered_results
    ]

    # Sort by similarity (best first) — Chroma doesn't guarantee order
    scored_results.sort(key=lambda x: x[1], reverse=True)

    # Drop low-quality chunks that add noise
    scored_results = [(doc, score) for doc, score in scored_results if score >= 0.35]

    # Limit to max results
    scored_results = scored_results[:MAX_RESULTS]

    return scored_results


def build_retriever(force_reload=False):
    """
    Build retriever for backward compatibility.
    Note: For accurate source filtering, use retrieve_with_scores() instead.
    """
    vectorstore = _ensure_vectorstore(force_reload)
    return vectorstore.as_retriever(
        search_type="similarity",
        search_kwargs={"k": TOP_K}
    )