"""
Canary watermarking service.

Embeds a unique tracking URL in each watermarked document.
When the document is opened by any browser, PDF viewer, or Office application,
the URL fires a GET request to /api/canary/ping/<token_hash> — capturing
the opener's IP address, User-Agent, and geolocation automatically.

Embedding strategy per format:
  PDF   — invisible URI annotation on page 1 (auto-fetched by most viewers)
  DOCX  — linked tracking image in a hidden paragraph
  TXT   — verification URL as plain text (fires if pasted into a browser)
"""
import os
import uuid
import hashlib
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument

OUTPUT_DIR = Path(__file__).resolve().parents[1] / "data" / "canary_output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Static canary keywords still used for in-text detection by the firewall
CANARY_TOKENS = ["zqxorin", "velmora", "kythrax"]
WATERMARK_TEXT = "Internal Reference: zqxorin velmora zqxorin"

# Base URL for the self-hosted ping endpoint.
# Override via CANARY_BASE_URL environment variable when running behind ngrok.
import os as _os
_DEFAULT_BASE = "http://localhost:5200"
CANARY_BASE_URL = _os.environ.get("CANARY_BASE_URL", _DEFAULT_BASE).rstrip("/")


def _ping_url(token_hash: str) -> str:
    return f"{CANARY_BASE_URL}/api/canary/ping/{token_hash}"


def _generate_token_hash() -> str:
    return hashlib.sha256(uuid.uuid4().bytes).hexdigest()[:32]


def generate_metadata(content: bytes) -> dict:
    canary_id  = str(uuid.uuid4())
    token_hash = _generate_token_hash()
    timestamp  = datetime.now(timezone.utc).isoformat()
    hash_val   = hashlib.sha256(content).hexdigest()
    return {
        "canary_id":   canary_id,
        "token_hash":  token_hash,
        "timestamp":   timestamp,
        "hash":        hash_val,
        "ping_url":    _ping_url(token_hash),
    }


def calculate_hash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


# ---------------------------------------------------------------------------
# HTML canary wrapper — fires in every browser, zero JavaScript required.
#
# When the recipient opens this .html file in Chrome / Edge / Firefox,
# the browser fetches the 1×1 <img> before rendering anything else.
# That GET hits /api/canary/ping/<hash> and we capture IP + UA + geo.
#
# The HTML also embeds the original document (base64) so the recipient
# still sees the actual content — it looks like a "secure document viewer".
# ---------------------------------------------------------------------------

def _make_html_canary(original_path: str, meta: dict, original_filename: str) -> str:
    """
    Wrap any document in an HTML page that fires a tracking pixel on open.
    Returns the path to the generated .html file.
    """
    ping          = meta["ping_url"]
    token_hash    = meta["token_hash"]
    canary_id     = meta["canary_id"]
    timestamp     = meta["timestamp"]
    doc_hash      = meta["hash"]

    # Read and base64-encode the original file for inline embedding
    import base64, mimetypes
    with open(original_path, "rb") as fh:
        raw = fh.read()
    b64 = base64.b64encode(raw).decode()
    mime = mimetypes.guess_type(original_filename)[0] or "application/octet-stream"
    data_uri = f"data:{mime};base64,{b64}"

    # Build the embedded viewer section depending on file type
    ext = Path(original_filename).suffix.lower()
    if ext == ".pdf":
        viewer_html = f'<embed src="{data_uri}" type="application/pdf" width="100%" height="100%" style="border:none;min-height:80vh;">'
    elif ext == ".txt":
        import html as _html
        text_content = _html.escape(raw.decode("utf-8", errors="replace"))
        viewer_html = f'<pre style="white-space:pre-wrap;word-break:break-word;background:#1e1e1e;color:#d4d4d4;padding:2rem;border-radius:8px;font-size:13px;line-height:1.6;overflow:auto;max-height:80vh;">{text_content}</pre>'
    else:
        # DOCX / unknown — show download link; can't render in browser natively
        viewer_html = f'''
        <div style="text-align:center;padding:4rem;">
          <p style="color:#94a3b8;font-size:1rem;margin-bottom:1.5rem;">This document requires Microsoft Word or a compatible viewer.</p>
          <a href="{data_uri}" download="{original_filename}"
             style="display:inline-block;padding:.75rem 2rem;background:#3b82f6;color:#fff;border-radius:8px;text-decoration:none;font-weight:600;">
            Download {original_filename}
          </a>
        </div>'''

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Secure Document — {original_filename}</title>
  <style>
    * {{ margin:0; padding:0; box-sizing:border-box; }}
    body {{ background:#0f172a; color:#f1f5f9; font-family:'Segoe UI',system-ui,sans-serif; min-height:100vh; display:flex; flex-direction:column; }}
    header {{ background:#1e293b; border-bottom:1px solid #334155; padding:.75rem 1.5rem; display:flex; align-items:center; gap:.75rem; }}
    header svg {{ width:20px; height:20px; fill:#3b82f6; flex-shrink:0; }}
    header .title {{ font-size:.95rem; font-weight:600; color:#f1f5f9; }}
    header .badge {{ margin-left:auto; font-size:.7rem; padding:.2rem .6rem; border-radius:999px; background:#0f172a; color:#94a3b8; border:1px solid #334155; font-family:monospace; }}
    main {{ flex:1; padding:1rem 1.5rem; }}
  </style>
</head>
<body>
  <!-- Tracking pixel — fires instantly on page load, no JS required -->
  <img src="{ping}" width="1" height="1"
       style="position:absolute;opacity:0;pointer-events:none;"
       alt="" referrerpolicy="unsafe-url">

  <header>
    <svg viewBox="0 0 24 24"><path d="M12 1L3 5v6c0 5.55 3.84 10.74 9 12 5.16-1.26 9-6.45 9-12V5l-9-4z"/></svg>
    <span class="title">ZeroSec — Secure Document Viewer</span>
    <span class="badge">ID: {canary_id[:8]}…</span>
  </header>

  <main>
    {viewer_html}
  </main>

  <!-- Canary metadata (visible in page source) -->
  <!-- Canary-ID: {canary_id} -->
  <!-- Token-Hash: {token_hash} -->
  <!-- SHA256: {doc_hash} -->
  <!-- Timestamp: {timestamp} -->
  <!-- Ping: {ping} -->
</body>
</html>"""

    orig_stem = Path(original_filename).stem
    output_path = OUTPUT_DIR / (orig_stem + "_canary.html")
    with open(output_path, "w", encoding="utf-8") as fh:
        fh.write(html)
    return str(output_path)


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------

def watermark_txt(input_path: str) -> dict:
    with open(input_path, "rb") as f:
        original_content = f.read()
    meta = generate_metadata(original_content)

    header = (
        f"{WATERMARK_TEXT}\n"
        f"Canary-ID: {meta['canary_id']}\n"
        f"Timestamp: {meta['timestamp']}\n"
        f"SHA256: {meta['hash']}\n"
        f"Verification: {meta['ping_url']}\n"
    ).encode("utf-8")

    orig_stem = Path(input_path).stem
    if orig_stem.startswith("temp_"):
        orig_stem = orig_stem[5:]
    output_path = OUTPUT_DIR / (orig_stem + "_canary.txt")
    with open(output_path, "wb") as f:
        f.write(header + original_content)
    meta["output_path"] = str(output_path)
    original_filename = orig_stem + ".txt"
    meta["html_path"] = _make_html_canary(str(output_path), meta, original_filename)
    return meta


# ---------------------------------------------------------------------------
# PDF — embeds an invisible URI action that fires when the page is rendered
# ---------------------------------------------------------------------------

def watermark_pdf(input_path: str) -> dict:
    with open(input_path, "rb") as f:
        original_content = f.read()
    meta = generate_metadata(original_content)
    ping = meta["ping_url"]

    try:
        from PyPDF2 import PdfReader, PdfWriter
        from PyPDF2.generic import (
            DictionaryObject, ArrayObject, NameObject,
            TextStringObject, NumberObject, DecodedStreamObject,
        )

        reader = PdfReader(input_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)

        # --- Strategy 1: /OpenAction with JavaScript ---
        # Fires automatically when Adobe Acrobat / Reader opens the document.
        # app.launchURL is a silent HTTP GET from Acrobat's perspective;
        # the user sees no dialog for external URLs in most Acrobat versions.
        js_code = f'app.launchURL("{ping}", true);'
        js_action = DictionaryObject({
            NameObject("/S"):  NameObject("/JavaScript"),
            NameObject("/JS"): TextStringObject(js_code),
        })
        writer._root_object.update({
            NameObject("/OpenAction"): js_action,
        })

        # --- Strategy 2: invisible URI Link annotation on page 1 ---
        # Fires when the user clicks anywhere near the (zero-size) hotspot,
        # or when crawlers/scanners follow embedded links.
        page0 = writer.pages[0]
        annot = DictionaryObject({
            NameObject("/Type"):    NameObject("/Annot"),
            NameObject("/Subtype"): NameObject("/Link"),
            NameObject("/Rect"):    ArrayObject([
                NumberObject(0), NumberObject(0),
                NumberObject(1), NumberObject(1),
            ]),
            NameObject("/F"):       NumberObject(2),   # hidden flag
            NameObject("/A"): DictionaryObject({
                NameObject("/S"):   NameObject("/URI"),
                NameObject("/URI"): TextStringObject(ping),
            }),
        })
        if "/Annots" not in page0:
            page0[NameObject("/Annots")] = ArrayObject()
        page0[NameObject("/Annots")].append(annot)

        # --- Strategy 3: metadata beacon ---
        # Token hash survives even if annotations are stripped.
        writer.add_metadata({
            "/Author":   f"ZeroSec Canary: {meta['canary_id']}",
            "/Subject":  f"Verification: {ping}",
            "/Keywords": f"canary:{meta['token_hash']}",
        })

        # Fix output filename — use original stem before the temp_ prefix
        orig_stem = Path(input_path).stem
        if orig_stem.startswith("temp_"):
            orig_stem = orig_stem[5:]   # strip "temp_" so output is clean
        output_path = OUTPUT_DIR / (orig_stem + "_canary.pdf")
        with open(output_path, "wb") as f:
            writer.write(f)

    except Exception:
        # Fallback: copy original
        import shutil
        orig_stem = Path(input_path).stem
        if orig_stem.startswith("temp_"):
            orig_stem = orig_stem[5:]
        output_path = OUTPUT_DIR / (orig_stem + "_canary.pdf")
        shutil.copy2(input_path, output_path)

    meta["output_path"] = str(output_path)
    original_filename = orig_stem + ".pdf"
    meta["html_path"] = _make_html_canary(str(output_path), meta, original_filename)
    return meta


# ---------------------------------------------------------------------------
# DOCX — adds a linked tracking image (Office auto-fetches linked images)
# ---------------------------------------------------------------------------

def watermark_docx(input_path: str) -> dict:
    doc = DocxDocument(input_path)
    with open(input_path, "rb") as f:
        original_content = f.read()
    meta = generate_metadata(original_content)
    ping = meta["ping_url"]

    # Add a white-on-white "verification" paragraph containing the URL.
    # Office and LibreOffice auto-fetch linked images; the URL in text also
    # shows up in metadata scanners.
    p = doc.add_paragraph()
    run = p.add_run(f"[Verification: {ping}]")
    # Make the text white (invisible against white background) using XML
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = run._r.get_or_add_rPr()
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FFFFFF')  # white text
    rPr.append(color)

    # Also add a linked-image relationship that points to the ping URL.
    # Word will attempt to load this when the document is opened online.
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        doc.part.relate_to(ping, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image',
                           is_external=True)
    except Exception:
        pass  # non-critical

    doc.add_paragraph(WATERMARK_TEXT)
    doc.add_paragraph(f"Canary-ID: {meta['canary_id']}")
    doc.add_paragraph(f"Timestamp: {meta['timestamp']}")
    doc.add_paragraph(f"SHA256: {meta['hash']}")

    orig_stem = Path(input_path).stem
    if orig_stem.startswith("temp_"):
        orig_stem = orig_stem[5:]
    output_path = OUTPUT_DIR / (orig_stem + "_canary.docx")
    doc.save(str(output_path))
    meta["output_path"] = str(output_path)
    original_filename = orig_stem + ".docx"
    meta["html_path"] = _make_html_canary(str(output_path), meta, original_filename)
    return meta
